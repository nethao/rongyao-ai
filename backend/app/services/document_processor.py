"""
文档处理服务
"""
import os
import subprocess
import tempfile
from pathlib import Path
from typing import Optional, List, Tuple
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import logging

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """文档处理器"""
    
    def __init__(self, temp_dir: Optional[str] = None):
        """
        初始化文档处理器
        
        Args:
            temp_dir: 临时文件目录
        """
        self.temp_dir = temp_dir or tempfile.gettempdir()
        os.makedirs(self.temp_dir, exist_ok=True)
    
    def convert_doc_to_docx(
        self,
        doc_path: str,
        output_dir: Optional[str] = None
    ) -> str:
        """
        使用LibreOffice将.doc转换为.docx
        
        Args:
            doc_path: .doc文件路径
            output_dir: 输出目录（默认为临时目录）
        
        Returns:
            str: 转换后的.docx文件路径
        
        Raises:
            RuntimeError: 转换失败时抛出
        """
        if not os.path.exists(doc_path):
            raise FileNotFoundError(f"文件不存在: {doc_path}")
        
        # 确定输出目录
        output_dir = output_dir or self.temp_dir
        os.makedirs(output_dir, exist_ok=True)
        
        try:
            # 调用LibreOffice进行转换
            # --headless: 无界面模式
            # --convert-to docx: 转换为docx格式
            # --outdir: 输出目录
            cmd = [
                'libreoffice',
                '--headless',
                '--convert-to', 'docx',
                '--outdir', output_dir,
                doc_path
            ]
            
            logger.info(f"开始转换文档: {doc_path}")
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=60  # 60秒超时
            )
            
            if result.returncode != 0:
                error_msg = result.stderr or result.stdout
                logger.error(f"LibreOffice转换失败: {error_msg}")
                raise RuntimeError(f"文档转换失败: {error_msg}")
            
            # 构建输出文件路径
            doc_filename = Path(doc_path).stem
            docx_path = os.path.join(output_dir, f"{doc_filename}.docx")
            
            if not os.path.exists(docx_path):
                raise RuntimeError(f"转换后的文件不存在: {docx_path}")
            
            logger.info(f"文档转换成功: {docx_path}")
            return docx_path
        
        except subprocess.TimeoutExpired:
            logger.error("LibreOffice转换超时")
            raise RuntimeError("文档转换超时")
        
        except Exception as e:
            logger.error(f"文档转换异常: {str(e)}")
            raise
    
    def extract_text_from_docx(self, docx_path: str, skip_title_lines: int = 0) -> str:
        """
        从.docx文件提取文本内容（带图片占位符）
        
        Args:
            docx_path: .docx文件路径
            skip_title_lines: 跳过前N行（用于跳过标题）
        
        Returns:
            str: 提取的文本内容，图片位置用[[IMG_N]]标记
        """
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"文件不存在: {docx_path}")
        
        try:
            doc = Document(docx_path)
            
            # 提取所有段落和表格的文本
            text_parts = []
            image_counter = 1
            line_counter = 0
            
            for element in doc.element.body:
                if isinstance(element, CT_P):
                    # 段落
                    paragraph = Paragraph(element, doc)
                    
                    # 检查段落中是否有图片
                    has_image = False
                    for run in paragraph.runs:
                        if run._element.xpath('.//a:blip'):
                            # 段落中有图片，插入占位符
                            text_parts.append(f'[[IMG_{image_counter}]]')
                            image_counter += 1
                            has_image = True
                            break
                    
                    # 如果没有图片，添加文本
                    if not has_image and paragraph.text.strip():
                        line_counter += 1
                        # 跳过标题行
                        if line_counter > skip_title_lines:
                            text_parts.append(paragraph.text)
                
                elif isinstance(element, CT_Tbl):
                    # 表格
                    table = Table(element, doc)
                    for row in table.rows:
                        row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                        if row_text.strip():
                            text_parts.append(row_text)
            
            content = '\n'.join(text_parts)
            logger.info(f"成功提取文本，长度: {len(content)}, 图片数量: {image_counter - 1}, 跳过行数: {skip_title_lines}")
            return content
        
        except Exception as e:
            logger.error(f"提取文本失败: {str(e)}")
            raise
    
    def extract_title_from_docx(self, docx_path: str) -> tuple[str, int]:
        """
        从.docx文件提取标题
        
        提取前两行非空内容作为标题：
        第1行：主标题
        第2行（如果以--开头）：副标题
        
        Args:
            docx_path: .docx文件路径
        
        Returns:
            tuple[str, int]: (提取的标题, 标题占用的行数)
        """
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"文件不存在: {docx_path}")
        
        try:
            doc = Document(docx_path)
            
            lines = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text and not text.startswith('[[IMG_'):
                    lines.append(text)
                    if len(lines) >= 2:
                        break
            
            if not lines:
                return "无标题", 0
            
            # 如果只有一行，直接返回
            if len(lines) == 1:
                return lines[0], 1
            
            # 如果第二行以--开头，组合为完整标题
            if lines[1].startswith('--'):
                subtitle = lines[1].lstrip('-').strip()
                return f"{lines[0]}——{subtitle}", 2
            else:
                # 否则只返回第一行
                return lines[0], 1
        
        except Exception as e:
            logger.error(f"提取标题失败: {str(e)}")
            return "无标题", 0
    
    def extract_images_from_docx(
        self,
        docx_path: str,
        output_dir: Optional[str] = None
    ) -> List[Tuple[str, bytes]]:
        """
        从.docx文件提取图片
        
        Args:
            docx_path: .docx文件路径
            output_dir: 图片输出目录（可选）
        
        Returns:
            List[Tuple[str, bytes]]: 图片列表，每项为(文件名, 图片数据)
        """
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"文件不存在: {docx_path}")
        
        try:
            doc = Document(docx_path)
            images = []
            
            # 遍历文档中的所有关系
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    # 获取图片数据
                    image_part = rel.target_part
                    image_data = image_part.blob
                    
                    # 获取图片扩展名
                    content_type = image_part.content_type
                    ext = content_type.split('/')[-1]
                    if ext == 'jpeg':
                        ext = 'jpg'
                    
                    # 生成文件名
                    filename = f"image_{len(images) + 1}.{ext}"
                    
                    images.append((filename, image_data))
                    
                    # 如果指定了输出目录，保存图片
                    if output_dir:
                        os.makedirs(output_dir, exist_ok=True)
                        image_path = os.path.join(output_dir, filename)
                        with open(image_path, 'wb') as f:
                            f.write(image_data)
            
            logger.info(f"成功提取 {len(images)} 张图片")
            return images
        
        except Exception as e:
            logger.error(f"提取图片失败: {str(e)}")
            raise
    
    def check_libreoffice_installed(self) -> bool:
        """
        检查LibreOffice是否已安装
        
        Returns:
            bool: 是否已安装
        """
        try:
            result = subprocess.run(
                ['libreoffice', '--version'],
                capture_output=True,
                text=True,
                timeout=5
            )
            return result.returncode == 0
        except (subprocess.TimeoutExpired, FileNotFoundError):
            return False
